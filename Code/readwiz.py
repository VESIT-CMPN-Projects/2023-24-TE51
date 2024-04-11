import streamlit as st
from langchain.text_splitter import CharacterTextSplitter
from langchain.document_loaders import PyPDFLoader, DirectoryLoader
from transformers import pipeline
from transformers import T5Tokenizer, T5ForConditionalGeneration
from gtts import gTTS
import io
import base64
import os
import torch
import tempfile
from streamlit_option_menu import option_menu
from langchain_google_genai import GoogleGenerativeAIEmbeddings 
import google.generativeai as genai 
from langchain.vectorstores import FAISS 
from PyPDF2 import PdfReader
from langchain.chains.question_answering import load_qa_chain 
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import ChatGoogleGenerativeAI
import os
import streamlit as st
from transformers import PegasusForConditionalGeneration, PegasusTokenizer
from sentence_splitter import SentenceSplitter, split_text_into_sentences
from fpdf import FPDF
from docx import Document
from PIL import Image
import win32com.client as win32
import matplotlib.pyplot as plt
from streamlit_markmap import markmap
from youtube_transcript_api import YouTubeTranscriptApi
from streamlit_player import st_player

print("Checking..")

offload_fld = "" #Enter your offload folder path here (if not created, create a folder and enter the path here)
refinedWiz = "RefinedWiz-Model" #Enter the model path here

tokenizer = T5Tokenizer.from_pretrained(refinedWiz)
base_model = T5ForConditionalGeneration.from_pretrained(refinedWiz, device_map="auto", torch_dtype=torch.float32, offload_folder=offload_fld)
para_model = 'tuner007/pegasus_paraphrase'
# torch_device = 'cuda' if torch.cuda.is_available() else 'cpu'
para_tokenizer = PegasusTokenizer.from_pretrained(para_model)
para_model = PegasusForConditionalGeneration.from_pretrained(para_model)

# Here we need to set the API key for the Gemini API, so get the API key from https://aistudio.google.com/app and then save the API key in the .env file with the key GOOGLE_API_KEY and then load the API key using the load_dotenv() function.
load_dotenv()
os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=os.getenv("GOOGLE_API_KEY")) 
text_model= genai.GenerativeModel("gemini-pro")
image_model = genai.GenerativeModel("gemini-pro-vision")


# functions for Summary
# process the file and get the text
def file_preprocess(file, selected_extension=None):
    if selected_extension == "pdf":
        loader = PyPDFLoader(file)
        pages = loader.load_and_split()
        text_splitter = CharacterTextSplitter(chunk_size=200, chunk_overlap=50)
        final_text = ""
        for text in pages:
            final_text += text.page_content
        # print(final_text)
        return final_text
    elif selected_extension == "txt":
        with open(file, "r", encoding="utf-8") as f:
            text = f.read()
            return text
    elif selected_extension == "docx":
        final_text = ""
        doc = Document(file)
        for para in doc.paragraphs:
            final_text += para.text
        return final_text
    elif selected_extension == "pptx":
        source = os.path.join(os.getcwd(), file)
        destination = os.path.join(os.getcwd(), file + ".pdf")
        print(destination)
        print(source)
        powerpoint = win32.Dispatch("Powerpoint.Application")
        deck = powerpoint.Presentations.Open(source)
        deck.SaveAs(destination, 32)
        deck.Close()
        powerpoint.Quit()
        # pass the file to file_preprocess
        print("going to ", destination)
        final_text =  file_preprocess(destination, "pdf")
        return final_text
    else:
        raise ValueError("Invalid file extension selected.")
    
    
# summarise the text
def llm_pipeline(input_text, option_selected, fileUpload = 0, selected_extension=None):
    if fileUpload == 1 and selected_extension == "pdf":
        input_text = file_preprocess(input_text, "pdf")
    elif fileUpload == 1 and selected_extension == "txt":
        input_text = file_preprocess(input_text, "txt")
    elif fileUpload == 1 and selected_extension == "docx":
        input_text = file_preprocess(input_text, "docx")
    elif fileUpload == 1 and selected_extension == "pptx":
        input_text = file_preprocess(input_text, "pptx")

    print("Input text: ", input_text)
    input_text = input_text.replace("\n", " ")

    splitter = SentenceSplitter(language='en')
    sentence_list = splitter.split(input_text)    
    join_n_sentence = []
    print("Sentence list: ", sentence_list)
    length = len(sentence_list)
    print(length)
    if length <= 3: 
        return input_text, input_text
    print("total sentences: ",length)
    if option_selected == 5:
        # 20% of the total sentences
        no_of_lines = (length * 60)//100
    elif option_selected == 10:
        # 40% of the total sentences
        no_of_lines = (length * 40)//100
    elif option_selected == 20:
        # 60% of the total sentences
        no_of_lines = (length * 20)//100
    else:
        raise ValueError("Invalid option selected.")
    print("No of lines: ", no_of_lines)
    for i in range(0, length, no_of_lines):
        join_n_sentence.append(" ".join(sentence_list[i:i+no_of_lines]))
    print("Join n sentence: ", join_n_sentence)
    print("Length of join n sentence: ", len(join_n_sentence))
    print("no of new lines: ", no_of_lines)
    # summarise each element of the list
    summary = []
    pipe_sum = pipeline("summarization", model=base_model, tokenizer=tokenizer)

    for i in join_n_sentence:
        summary.append(pipe_sum(i))
    final_summary = ""
    for i in summary:
        final_summary += i[0]['summary_text'] + " "
    return final_summary, input_text

# complete the sentence if not ended with punctuation
def complete_sentence(summary, original_text):
    if not summary.endswith((".", "!", "?")):
        last_sentence = original_text.rsplit(".", 2)[-2] 
        last_sentence = last_sentence.strip()
        complete_summary = summary + ". " + last_sentence + "."
        return complete_summary
    else:
        return summary

@st.cache_data
# display the PDF file or the text
def displayFile(file, selected_extension=None):
    if selected_extension == "pdf":
        with open(file, "rb") as f:
            base64_pdf = base64.b64encode(f.read()).decode('utf-8')
        pdf_disp = f'<embed src="data:application/pdf;base64,{base64_pdf}" width="100%" height="700" type="application/pdf">'
        st.markdown(pdf_disp, unsafe_allow_html=True)
    elif selected_extension == "txt":
        with open(file, "r", encoding="utf-8") as f:
            text = f.read()
            st.warning(text)
    elif selected_extension == "docx":
        doc = Document(file)
        for para in doc.paragraphs:
            st.warning(para.text)
    elif selected_extension == "pptx":
        source = os.path.join(os.getcwd(), file)
        destination = os.path.join(os.getcwd(), file + ".pdf")
        print(destination)
        print(source)
        powerpoint = win32.Dispatch("Powerpoint.Application")
        deck = powerpoint.Presentations.Open(source)
        deck.SaveAs(destination, 32)
        deck.Close()
        powerpoint.Quit()
        with open(destination, "rb") as f:
            base64_pdf = base64.b64encode(f.read()).decode('utf-8')
        pdf_disp = f'<embed src="data:application/pdf;base64,{base64_pdf}" width="100%" height="700" type="application/pdf">'
        st.markdown(pdf_disp, unsafe_allow_html=True)
    else:
        raise ValueError("Invalid file extension selected.")


# functions for Chat
# function for question answering with PDF
def get_answer(question, input_text):
    if question== "":
        return "Please enter a question."
    if input_text == "":
        return "Please upload a PDF."
    text = "You are asked a question which is: " + question + " and you have to generate an answer based on the given text: " + input_text + ". IT should not be complicated and easy to understand. The answer should be from the input text given to you only, if there is something asked outside from the input_text, then tell the user that the answer is not present in the given text. Tho give the user some kindoff small answer mentioning its not given its from your own. If there is something important asked from the PDF and there is no answer to that, dont answer it then. Dont answer in brief unless asked to."
    response = text_model.generate_content(text)
    response.resolve()
    return response.text

# function for paraphrase
# function for paraphrasing the text
def get_response(input_text,num_return_sequences):
    input_text = input_text.replace("\n", " ")
    inputs = para_tokenizer([input_text], return_tensors="pt", padding=True, truncation=True)
    summary_ids = para_model.generate(inputs.input_ids, num_return_sequences=num_return_sequences, num_beams=10, no_repeat_ngram_size=2, top_k=50, top_p=0.95, temperature=0.7, do_sample=True, max_length=100)
    paraphrase = para_tokenizer.batch_decode(summary_ids, skip_special_tokens=True)
    return paraphrase

# function for getting the paraphrased text
def get_paraphrase(input_text, num_sentences):
    input_text = input_text.replace("\n", " ")
    splitter = SentenceSplitter(language='en')
    sentence_list = splitter.split(input_text)
    # print(input_text)
    paraphrase = []
    for sentence in sentence_list:
        paraphrase.append(get_response(sentence,num_sentences))
    all_paraphrases = [] # list of all lists of paraphrased sentences
    for i in range(num_sentences):
        paraphrases = []
        for j in paraphrase:
            paraphrases.append(j[i])
        all_paraphrases.append(paraphrases)
    paraphrased_sentences = [] # list of paraphrased sentences into a each different string
    for i in all_paraphrases:
        para = ""
        for j in i:
            para = para + j + " "
        paraphrased_sentences.append(para)
    final_list = "" # all list into a single string with new line
    for i in paraphrased_sentences:
        final_list += i + "\n\n------------------------------------------------------------\n\n"
    print(final_list)
    return final_list

# function for text to speech
def speak_text(text):
    # language = st.selectbox("Select Language", ("en", "hi", "es", "fr", "de", "it", "ja", "ko", "zh-CN", "zh-TW"))
    try:
        tts = gTTS(text, lang="en")
        tts.save("temp/temp.mp3")
        audio_file = open("temp/temp.mp3", "rb")
        audio_bytes = audio_file.read()
        st.audio(audio_bytes, format="audio/ogg")
    except:
        st.write("Please try again")

# function for analyzing the summary and displaying the results
def analyze_summary(input_text, summ):
    original_size = len(input_text.split())
    summary_size = len(summ.split())
    original_sentences = len(split_text_into_sentences(input_text, language='en'))
    summary_sentences = len(split_text_into_sentences(summ, language='en'))
    original_characters = len(input_text)
    summary_characters = len(summ)
    reduction_in_size = original_size - summary_size
    reduction_in_sentences = original_sentences - summary_sentences
    reduction_in_characters = original_characters - summary_characters
    reduced_percent = (reduction_in_size/original_size) * 100
    col1, col2 = st.columns(2)
    with col1:
        st.write(f"Original Size: {original_size} words")
        st.write(f"Original Sentences: {original_sentences}")
        st.write(f"Original Characters: {original_characters}")
    with col2:
        st.write(f"Summary Size: {summary_size} words")
        st.write(f"Summary Sentences: {summary_sentences}")
        st.write(f"Summary Characters: {summary_characters}")
    
    st.markdown("---")
    col1, col2 = st.columns(2)
    with col1:
        
        labels = 'Remaining total', 'Reduced Summary', 
        sizes = [reduction_in_size, summary_size]
        fig1, ax1 = plt.subplots()
        ax1.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90, radius=0.1, colors=['#ff9999','#66b3ff'])
        ax1.axis('equal')
        st.pyplot(fig1)
    with col2:
        if reduction_in_sentences > original_sentences:
            st.info(f'''Reduced by: {reduction_in_size} words & {reduction_in_characters} characters.
                
                Reduced in total size by: {reduced_percent:.2f}%''')
        else:
            st.info(f'''Reduced by: {reduction_in_size} words, {reduction_in_sentences} sentences, {reduction_in_characters} characters.
                
                Reduced in total size by: {reduced_percent:.2f}%''')

# function for summarizing the image
def get_img_summary(uploaded_file):
    min_length = 20
    text = f"Summarize the given image, like identify the objects, people, animals, etc. and give a brief summary of the image. The summary should be easy to understand and should not be too long. The summary should be of minimum 20 words. Dont overdo it, just give a brief summary of the image."
    response = image_model.generate_content([text, image])
    return response.text

# function for getting the mindmap md content
def generate_markdown(text):
    query = rf"""
        Study the given {text} and generate a summary then please be precise in selecting the data such that it gets to a heirarchical structure. Dont give anything else, i just want to display the structure as a mindmap so be precise please. Dont write anything else, Just return the md file. It is not neccessay to cover all information. dont use triple backticks or ` anywhere. Cover the main topics. Please convert this data into a markdown mindmap format similar to the following example:
        ---
        markmap:
        colorFreezeLevel: 2
        ---

        # Gemini Account Summary

        ## Balances

        - Bitcoin (BTC): 0.1234
        - Ethereum (ETH): 0.5678

        ## Orders

        - Open Orders
        - Buy Order (BTC): 0.01 BTC @ $40,000
        - Trade History
        - Sold 0.1 ETH for USD at $2,500

        ## Resources

        - [Gemini Website](https://www.gemini.com/)
    """
    response = text_model.generate_content(query)
    response.resolve()
    # save the response to a markdown file
    with open("summary.md", "w") as f:
        f.write(response.text)
    return response.text

# function to display mindmap
def generate_mindmap(mindmap_summ):
    mindmap = markmap(mindmap_summ)
    return mindmap

# function to get all the difficult words from the text
def get_difficult_words(text):
    query = f"""
    Study the given {text} and return the difficult words that are hard to understand for a 10 year old child. Dont change the words, just return the words that are hard to understand. get most of the difficult words from the content.  just return the words not a line or sentence or a group of 2-3 words. just return the words that are hard to understand.
    """
    response = text_model.generate_content(query)
    response.resolve()
    print(response.text)
    return response.text

# function to get the meaning of the difficult words
def generate_para(text, difficult_words):
    query = f"""
    Study the given {text} properly and return the whole paragraph in <p> tag of html, but add the meaning of every {difficult_words} that are present in the paragraph. right after the word. just give the meaning of word not the line or sentence, dont give meaning of useless words check the content and findout the difficult words only that are hard to understand. make sure the meaning is simple to understand, the meaning should be in the <span> tag. with the class name as "meaning" and data-hover as the meaning of the word. Make sure the meaning is different from the word itself and simple to understand for the child. return the paragraph as it was as {text}.
    example:
    <p>
        The <span class="meaning" data-hover="big">large</span> brown fox jumped over the lazy dog.
    </p> 
    """
    print("hogaya query")
    response = text_model.generate_content(query)
    response.resolve()
    return response.text

# function to copy the text 
def copy_to_clip(text):
    query = f"""
    Study the given {text}, return thw whole paragraph in <p> tag of html and class should be "paragraph".
    like this:
    <p class="paragraph">{text}</p>
    """
    response = text_model.generate_content(query)
    response.resolve()
    return response.text

# Set the page configuration
st.set_page_config(layout="wide", page_title="ReadWiz - Text Summarizer & Paraphraser", page_icon=":books:")
# SIDEBAR
with st.sidebar:
    st.title("ReadWiz")
    selected = option_menu("",['Get Summary', 'Get Paraphrase', 'Chat with PDF', 'Image || Video Summary',   'About'], 
        icons=['text-center', 'text-paragraph', 'question', 'image', 'people'], menu_icon="cast", default_index=0)
    st.markdown("---")
    st.write("Developed by Group 51")
    st.markdown("---")


# SUMMARY PAGE
if selected == "Get Summary":
    # st.write("Summary for the text")
    st.title("ReadWiz - Text Summarization📝")
    text_input = st.text_area("Enter Text:", height=60)
    st.title("OR")
    # uploaded_file = st.file_uploader("Upload PDF File", type=['pdf'])
    file_path = None
    colx, coly = st.columns(2)
    with colx:
        extensions = {
            "pdf": "pdf",
            "txt": "txt",
            "docx": "docx",
            "pptx": "pptx",
        }
        selected_extension = st.selectbox(  
            "Choose the file extension:",
            options=list(extensions.keys()),
        )
    with coly:
        uploaded_file = st.file_uploader("Upload File", type=selected_extension)
    # slider
    slider_options = {
        "Short": 5,  # Adjust values to model's constraints
        "Medium": 10,
        "Long": 20,
    }
    selected_option = st.selectbox(
        " Choose the summary length:",
        options=list(slider_options.keys()),
    )
    option_selected = slider_options[selected_option]
    print(selected_extension)
    print(option_selected) # 5, 10, 20
    if st.button("Summarize"):
        if text_input or uploaded_file:
            col1, col2 = st.columns(2)
            input_text = text_input
            if uploaded_file:
                file_path = os.path.join(tempfile.gettempdir(), uploaded_file.name)
                print(file_path)
                # uploaded_file = check_extension(uploaded_file, selected_extension)
                with open(file_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
                input_text = file_preprocess(file_path, selected_extension)
                print(input_text)
            with col1:
                if uploaded_file:
                    st.info("Uploaded FILE")
                    pdf_viewer = displayFile(file_path, selected_extension)
                else:
                    st.info("Original Text")
                    st.warning(input_text)
            with col2:
                st.info("Summarized Text")
                if uploaded_file:
                    summ, input_text = llm_pipeline(file_path, option_selected, 1, selected_extension)
                else:
                    summ, input_text = llm_pipeline(input_text, option_selected,0, selected_extension)
                # final = complete_sentence(summ, input_text)  # Pass original text for completion
                # st.success(final)        
                difficult_words = get_difficult_words(summ)   
                content = generate_para(summ, difficult_words)    
                # st.success(summ)
                print(summ)
                st.markdown("""
                <style>
                    .meaning {
                    position: relative;
                    border-bottom: 1px dotted black;
                    /*color: #21244c;*/
                    }

                    .meaning:before {
                    content: attr(data-hover);
                    visibility: hidden;
                    opacity: 0;
                    width: 200px;
                    background-color: #8187c9;
                    text-align: center;
                    border-radius: 5px;
                    padding: 5px 0;
                    transition: opacity 0.5s ease-in-out;
                    color: white;

                    position: absolute;
                    z-index: 1;
                    left: 0;
                    top: 110%;
                    }

                    .meaning:hover:before {
                    opacity: 1;
                    visibility: visible;
                    }      
                </style>
                """, unsafe_allow_html=True)

                st.markdown(content, unsafe_allow_html=True)
                board_text = copy_to_clip(summ)
                # st.warning("Click this text to copy to clipboard")
                styling = """
                <style>
                .copy-board {
                    
                    position: relative;
                    border-bottom: 1px dotted black;
                    cursor: pointer;
                }
                .copy-board:before {
                    content: attr(data-hover);
                    visibility: hidden;
                    opacity: 0;
                    width: 200px;
                    background-color: #8187c9;
                    color: white;
                    text-align: center;
                    border-radius: 5px;
                    padding: 5px 0px;
                    transition: opacity 0.4s ease-in-out;

                    position: absolute;
                    z-index: 1;
                    left: 0;
                    top: 110%;
                }

                .copy-board:hover:before {
                    opacity: 1;
                    visibility: visible;
                }
                </style>    
                """
                    
                copy_script = f"""
                <p class="copy-board" data-hover="Copy">📄</p>
                <script>
                    document.addEventListener('click', function(e) {{
                        if (e.target.classList.contains('paragraph')) {{
                            var dummy = document.createElement("textarea");
                            document.body.appendChild(dummy);
                            dummy.value = "{summ}";
                            dummy.select();
                            document.execCommand("copy");
                            document.body.removeChild(dummy);
                        }}
                    }});
                </script>
                """
                # st.markdown(board_text, unsafe_allow_html=True)
                col1, col2 = st.columns(2)
                with col1:
                    st.success("Speech for the Summary:")
                with col2:
                    speak_text(summ)
                # with col3:
                #     st.markdown(styling, unsafe_allow_html=True)
                #     st.markdown(copy_script, unsafe_allow_html=True)

            st.markdown("---")
            st.success("MindMap of the Summary::")
            mindmap_summ = generate_markdown(summ)
            mindmap = generate_mindmap(mindmap_summ)
            # st.markdown(mindmap, unsafe_allow_html=True)

            st.markdown("---")
            st.success("Analysis of the Summary::")
            analyze_summary(input_text, summ)
            
        else:
            st.error("Please enter text or upload a PDF file.")



# PARAPHRASE
elif selected == "Get Paraphrase":
    # st.write("Paraphrase the text")
    st.title("ReadWiz - Text Paraphraser📝")
    text_input = st.text_area("Enter Text:", height=60)
    st.title("OR")
    # uploaded_file = st.file_uploader("Upload PDF File", type=['pdf'])
    file_path = None
    colx, coly = st.columns(2)
    with colx:
        extensions = {
            "pdf": "pdf",
            "txt": "txt",
            "docx": "docx",
            "pptx": "pptx",
            "image": "image",
        }
        selected_extension = st.selectbox(  
            "Choose the file extension:",
            options=list(extensions.keys()),
        )
    with coly:
        if selected_extension == "image":
            uploaded_file = st.file_uploader("Upload Image File", type=["png", "jpg", "jpeg"])
            st.write("Soon to be implemented")
        else:
            uploaded_file = st.file_uploader("Upload File", type=selected_extension)
    if st.button("Paraphrase"):
        if text_input or uploaded_file:
            col1, col2 = st.columns(2)
            input_text = text_input
            if uploaded_file:
                file_path = os.path.join(tempfile.gettempdir(), uploaded_file.name)
                with open(file_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
                input_text = file_preprocess(file_path, selected_extension)
            with col1:
                if uploaded_file:
                    st.info("Uploaded FILE")
                    pdf_viewer = displayFile(file_path, selected_extension)
                else:
                    st.info("Original Text")
                    st.warning(input_text)
            with col2:
                st.info("Paraphrased Text")
                # summ = get_paraphrase(input_text, 1)
                if uploaded_file:
                    # display all the items of list on new line
                    summ = get_paraphrase(input_text, 3)
                else:
                    summ = get_paraphrase(input_text, 3)
                st.success(summ)               
        else:
            st.error("Please enter text or upload a PDF file.")


# CHAT WITH PDF
elif selected == "Chat with PDF":
    # st.write("AI Chatbot with PDF")
    st.title("ReadWiz - Chat with PDF🗣️")
    uploaded_file = st.file_uploader("Upload PDF File", type=['pdf'])
    file_path = None
    input_text = ""
    if uploaded_file:
        file_path = os.path.join(tempfile.gettempdir(), uploaded_file.name)
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        input_text = file_preprocess(file_path, "pdf")
        st.info("Uploaded PDF")
        pdf_viewer = displayFile(file_path, "pdf")
        st.markdown("---")
        question = st.text_input("Enter your question here:")
        if st.button("Ask"):
            answer = get_answer(question, input_text)
            st.success(answer)
    else:
        st.error("Please upload a PDF file.")



# IMAGE OR VIDEO SUMMARY
elif selected == "Image || Video Summary":
    st.title("ReadWiz - Image || Video Summary📝")
    # st.write("Soon to be implemented")
    st.write("Please upload an image or enter youtube video link to get the summary.")
    uploaded_file = st.file_uploader("Upload Image File", type=["png", "jpg", "jpeg"])
    st.subheader("OR")
    video_link = st.text_input("Enter the Youtube Video Link:")
    # if uploaded_file:
    #     # st.image(uploaded_file, caption="Uploaded Image", use_column_width=True)
    #     max_length = st.slider("Select length of the summary", 50, 100, 70)
    # elif video_link:
    slider_options = {
        "Short": 5, 
        "Medium": 10,
        "Long": 20,
    }
    selected_option = st.selectbox(
        " Choose the summary length:",
        options=list(slider_options.keys()),
    )
    option_selected = slider_options[selected_option]
    print(option_selected) # 5, 10, 20
        
    if st.button("Get Summary"):
        if uploaded_file:
            col1, col2 = st.columns(2)
            with col1:
                st.info("Uploaded Image")
                image = Image.open(uploaded_file)
                st.image(image, caption="Uploaded Image", use_column_width=True)
            with col2:
                st.info("Summarized Text")
                
                summ = get_img_summary(uploaded_file)
                st.success(summ)
        elif video_link:
            if "www.youtube.com/watch?" in video_link:
                st_player(video_link)
                video_id = video_link.split("v=")[1]
                video_id = video_id.split("&")[0]
                try:
                    transcript = YouTubeTranscriptApi.get_transcript(video_id)
                    input_text = ""
                    for i in transcript:
                        input_text += i["text"] + " "
                    # remove [Music] and [Applause] from the text
                    input_text = input_text.replace("[Music]", "")
                    input_text = input_text.replace("[Applause]", "")
                    if not input_text.endswith((".", "!", "?")):
                        input_text = input_text+ "."
                    input_text = input_text.replace(" i ", ". I ")
                    input_text = input_text.replace(" I ", ". I ")
                    summ, input_text = llm_pipeline(input_text, option_selected,0)
                    difficult_words = get_difficult_words(summ)   
                    content = generate_para(summ, difficult_words)    
                    # st.success(summ)
                    print(summ)
                    st.markdown("""
                    <style>
                        .meaning {
                        position: relative;
                        border-bottom: 1px dotted black;
                        color: blue;
                        }

                        .meaning:before {
                        content: attr(data-hover);
                        visibility: hidden;
                        opacity: 0;
                        width: 200px;
                        background-color: white;
                        color: black;
                        text-align: center;
                        border-radius: 5px;
                        padding: 5px 0;
                        transition: opacity 1s ease-in-out;

                        position: absolute;
                        z-index: 1;
                        left: 0;
                        top: 110%;
                        }

                        .meaning:hover:before {
                        opacity: 1;
                        visibility: visible;
                        }      
                    </style>
                    """, unsafe_allow_html=True)

                    st.markdown(content, unsafe_allow_html=True)
                    col1, col2 = st.columns(2)
                    with col1:
                        st.success("Speech for the Summary:")
                    with col2:
                        speak_text(summ)

                    st.markdown("---")
                    st.success("MindMap of the Summary::")
                    mindmap_summ = generate_markdown(summ)
                    mindmap = generate_mindmap(mindmap_summ)
                    # st.markdown(mindmap, unsafe_allow_html=True)

                    st.markdown("---")
                    st.success("Analysis of the Summary::")
                    if len(input_text) != len(summ):

                        analyze_summary(input_text, summ)
                    else:
                        # st.error("The summary is same as the original text.")
                        st.warning("The summary is same as the original text.")
                except:
                    st.error("Looks like there are no transcripts found for this video!.")
                    st.warning("Please try another video.")
                    
                
            else:
                st.error("Please enter a valid youtube video link.")
        else:
            st.error("Please upload an image or enter a youtube video link.")
    

# https://github.com/ikitcheng/chinamatt_youtube/blob/main/2020-02-14-Youtube_Transcript/get_youtube_transcript.py
        # https://www.youtube.com/watch?v=n8s9DjPDBEw


# MULTIPLE FILES
elif selected == "Multiple files":
    st.title("Soon to be implemented")
    
    

# ABOUT PAGE
elif selected == "About":
    st.title("About ReadWiz")
    st.write("ReadWiz is a text summarizer, paraphraser and chatbot application. \n\nIt is designed to help users to summarize, paraphrase and chat with the given text or PDF file. The application is built using Streamlit, HuggingFace's Transformers library, Google's GenerativeAI, and other Python libraries.")
    st.write("The application is developed by Group 51, which consists of the following members:")
    st.write("1. Manraj Singh Virdi (D12B - 66)")
    st.write("2. Chirag Santwani (D12B - 50)")
    st.write("3. Nikhil Dhanwani (D12B - 12)")
    st.markdown("---")
    st.write("Thank you for using the Project!")
    